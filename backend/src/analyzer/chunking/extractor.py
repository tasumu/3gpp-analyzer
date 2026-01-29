"""Document structure extractor using python-docx."""

import re
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document as DocxDocument
from docx.document import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from analyzer.models.chunk import StructureType


@dataclass
class StructureElement:
    """Represents a structural element extracted from a document."""

    content: str
    structure_type: StructureType
    heading_level: int | None = None
    clause_number: str | None = None
    page_number: int | None = None
    parent_headings: list[str] = field(default_factory=list)


class DocxExtractor:
    """
    Extracts document structure from docx files.

    Identifies:
    - Headings (by style or formatting)
    - Paragraphs
    - Lists
    - Tables
    """

    # Patterns for detecting clause numbers
    CLAUSE_PATTERN = re.compile(r"^(\d+(?:\.\d+)*)\s+")

    # Heading style names
    HEADING_STYLES = {
        "Heading 1": (StructureType.HEADING1, 1),
        "Heading 2": (StructureType.HEADING2, 2),
        "Heading 3": (StructureType.HEADING3, 3),
        "Heading 4": (StructureType.HEADING4, 4),
        "Heading 5": (StructureType.HEADING5, 5),
        "Heading 6": (StructureType.HEADING6, 6),
        "Title": (StructureType.TITLE, 0),
    }

    def __init__(self):
        """Initialize the extractor."""
        self._current_headings: dict[int, str] = {}

    def extract_structure(self, file_path: Path | str) -> list[StructureElement]:
        """
        Extract structural elements from a docx file.

        Args:
            file_path: Path to the docx file.

        Returns:
            List of StructureElement objects in document order.
        """
        file_path = Path(file_path)
        doc = DocxDocument(file_path)

        elements = []
        self._current_headings = {}

        for element in self._iter_block_items(doc):
            if isinstance(element, Paragraph):
                extracted = self._extract_paragraph(element)
                if extracted and extracted.content.strip():
                    elements.append(extracted)

            elif isinstance(element, Table):
                extracted = self._extract_table(element)
                if extracted and extracted.content.strip():
                    elements.append(extracted)

        return elements

    def _iter_block_items(self, doc: Document):
        """
        Iterate through all block-level items in document order.

        This handles the fact that tables and paragraphs are
        separate in the docx structure.
        """
        parent = doc.element.body
        for child in parent.iterchildren():
            if child.tag == qn("w:p"):
                yield Paragraph(child, doc)
            elif child.tag == qn("w:tbl"):
                yield Table(child, doc)

    def _extract_paragraph(self, para: Paragraph) -> StructureElement | None:
        """Extract structure from a paragraph."""
        text = para.text.strip()
        if not text:
            return None

        # Determine structure type from style
        style_name = para.style.name if para.style else ""
        structure_type = StructureType.PARAGRAPH
        heading_level = None

        # Check for heading styles
        for style, (stype, level) in self.HEADING_STYLES.items():
            if style.lower() in style_name.lower():
                structure_type = stype
                heading_level = level
                break

        # Check for list items
        if para._element.pPr is not None:
            num_pr = para._element.pPr.find(qn("w:numPr"))
            if num_pr is not None:
                structure_type = StructureType.LIST_ITEM

        # Extract clause number
        clause_number = None
        match = self.CLAUSE_PATTERN.match(text)
        if match:
            clause_number = match.group(1)

        # Update heading hierarchy
        if heading_level is not None:
            self._current_headings[heading_level] = text
            # Clear lower-level headings
            for level in list(self._current_headings.keys()):
                if level > heading_level:
                    del self._current_headings[level]

        # Get parent headings (excluding current)
        parent_headings = []
        for level in sorted(self._current_headings.keys()):
            if heading_level is None or level < heading_level:
                parent_headings.append(self._current_headings[level])

        return StructureElement(
            content=text,
            structure_type=structure_type,
            heading_level=heading_level,
            clause_number=clause_number,
            parent_headings=parent_headings,
        )

    def _extract_table(self, table: Table) -> StructureElement | None:
        """Extract structure from a table."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))

        content = "\n".join(rows)
        if not content.strip():
            return None

        # Get current heading hierarchy
        parent_headings = [
            self._current_headings[level]
            for level in sorted(self._current_headings.keys())
        ]

        return StructureElement(
            content=content,
            structure_type=StructureType.TABLE,
            parent_headings=parent_headings,
        )

    def extract_title(self, file_path: Path | str) -> str | None:
        """
        Extract the document title.

        Args:
            file_path: Path to the docx file.

        Returns:
            Document title if found.
        """
        file_path = Path(file_path)
        doc = DocxDocument(file_path)

        # Check core properties
        if doc.core_properties.title:
            return doc.core_properties.title

        # Look for Title style or first heading
        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ""
            if "title" in style_name.lower() or "heading 1" in style_name.lower():
                if para.text.strip():
                    return para.text.strip()

        return None
