"""Service for managing user-uploaded attachments."""

import io
import logging
import uuid
from pathlib import Path

from analyzer.models.attachment import Attachment
from analyzer.providers.firestore_client import FirestoreClient
from analyzer.providers.storage_client import StorageClient

logger = logging.getLogger(__name__)

ATTACHMENTS_COLLECTION = "qa_attachments"
ATTACHMENTS_GCS_PREFIX = "attachments"
MAX_FILE_SIZE = 20 * 1024 * 1024  # 20MB
ALLOWED_EXTENSIONS = {".docx", ".xlsx", ".xls", ".pdf", ".txt", ".csv"}


class AttachmentService:
    """Service for uploading, extracting text from, and managing file attachments."""

    def __init__(self, firestore: FirestoreClient, storage: StorageClient):
        self.firestore = firestore
        self.storage = storage

    async def upload(
        self,
        meeting_id: str,
        filename: str,
        content: bytes,
        content_type: str,
        uploaded_by: str,
        session_id: str | None = None,
    ) -> Attachment:
        """
        Upload a file, extract text, and store metadata.

        Args:
            meeting_id: Associated meeting ID.
            filename: Original filename.
            content: File content as bytes.
            content_type: MIME type.
            uploaded_by: User ID of uploader.

        Returns:
            Created Attachment record.

        Raises:
            ValueError: If file extension is not supported or file is too large.
        """
        # Validate file extension
        ext = Path(filename).suffix.lower()
        if ext not in ALLOWED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file extension: {ext}. "
                f"Allowed: {', '.join(sorted(ALLOWED_EXTENSIONS))}"
            )

        # Validate file size
        if len(content) > MAX_FILE_SIZE:
            raise ValueError(f"File too large: {len(content)} bytes (max {MAX_FILE_SIZE})")

        attachment_id = str(uuid.uuid4())

        # Upload original file to GCS
        gcs_path = f"{ATTACHMENTS_GCS_PREFIX}/{meeting_id}/{attachment_id}/{filename}"
        await self.storage.upload_bytes(content, gcs_path, content_type)

        # Extract text
        extracted_text = self._extract_text(filename, content)

        # Upload extracted text to GCS
        text_gcs_path = f"{ATTACHMENTS_GCS_PREFIX}/{meeting_id}/{attachment_id}/extracted.txt"
        await self.storage.upload_bytes(
            extracted_text.encode("utf-8"), text_gcs_path, "text/plain; charset=utf-8"
        )

        # Create Firestore record
        attachment = Attachment(
            id=attachment_id,
            filename=filename,
            content_type=content_type,
            meeting_id=meeting_id,
            gcs_path=gcs_path,
            extracted_text_gcs_path=text_gcs_path,
            file_size_bytes=len(content),
            uploaded_by=uploaded_by,
            session_id=session_id,
        )
        self.firestore.client.collection(ATTACHMENTS_COLLECTION).document(attachment_id).set(
            attachment.to_firestore()
        )

        logger.info(
            f"Uploaded attachment {attachment_id}: {filename} "
            f"({len(content)} bytes, {len(extracted_text)} chars extracted)"
        )
        return attachment

    def _extract_text(self, filename: str, content: bytes) -> str:
        """Extract text content based on file extension."""
        lower = filename.lower()
        if lower.endswith(".docx"):
            return self._extract_docx(content)
        elif lower.endswith((".xlsx", ".xls")):
            return self._extract_xlsx(content)
        elif lower.endswith((".txt", ".csv")):
            return content.decode("utf-8", errors="replace")
        else:
            return f"[Text extraction not supported for {filename}]"

    def _extract_docx(self, content: bytes) -> str:
        """Extract text from .docx using python-docx."""
        from docx import Document as DocxDocument

        doc = DocxDocument(io.BytesIO(content))
        parts: list[str] = []

        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        for table in doc.tables:
            # Convert table to markdown format
            rows_text: list[str] = []
            for i, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                rows_text.append("| " + " | ".join(cells) + " |")
                if i == 0:
                    rows_text.append("| " + " | ".join(["---"] * len(cells)) + " |")
            if rows_text:
                parts.append("\n".join(rows_text))

        return "\n\n".join(parts)

    def _extract_xlsx(self, content: bytes) -> str:
        """Extract text from .xlsx using openpyxl, format as markdown tables."""
        import openpyxl

        wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        sections: list[str] = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                continue

            section_parts: list[str] = [f"## Sheet: {sheet_name}\n"]

            # Header row
            header = rows[0]
            header_strs = [str(c) if c is not None else "" for c in header]
            section_parts.append("| " + " | ".join(header_strs) + " |")
            section_parts.append("| " + " | ".join(["---"] * len(header)) + " |")

            # Data rows
            for row in rows[1:]:
                cells = [str(c) if c is not None else "" for c in row]
                section_parts.append("| " + " | ".join(cells) + " |")

            sections.append("\n".join(section_parts))

        wb.close()
        return "\n\n".join(sections)

    async def list_by_meeting(
        self, meeting_id: str, session_id: str | None = None
    ) -> list[Attachment]:
        """List all attachments for a meeting, optionally filtered by session."""
        query = self.firestore.client.collection(ATTACHMENTS_COLLECTION).where(
            "meeting_id", "==", meeting_id
        )
        if session_id:
            query = query.where("session_id", "==", session_id)
        query = query.order_by("created_at", direction="DESCENDING")
        results = []
        for doc in query.stream():
            results.append(Attachment.from_firestore(doc.id, doc.to_dict()))
        return results

    async def get(self, attachment_id: str) -> Attachment | None:
        """Get a single attachment by ID."""
        doc_ref = self.firestore.client.collection(ATTACHMENTS_COLLECTION).document(attachment_id)
        doc = doc_ref.get()
        if doc.exists:
            return Attachment.from_firestore(doc.id, doc.to_dict())
        return None

    async def get_extracted_text_with_metadata(
        self, attachment_id: str
    ) -> tuple[Attachment, str] | tuple[None, None]:
        """Get attachment metadata and extracted text in a single Firestore read."""
        attachment = await self.get(attachment_id)
        if not attachment or not attachment.extracted_text_gcs_path:
            return None, None
        content = await self.storage.download_bytes(attachment.extracted_text_gcs_path)
        return attachment, content.decode("utf-8")

    async def get_extracted_text(self, attachment_id: str) -> str | None:
        """Get extracted text content of an attachment."""
        _, text = await self.get_extracted_text_with_metadata(attachment_id)
        return text

    async def delete(self, attachment_id: str, user_id: str) -> bool:
        """
        Delete an attachment (only by uploader).

        Args:
            attachment_id: Attachment to delete.
            user_id: User requesting deletion.

        Returns:
            True if deleted successfully.

        Raises:
            PermissionError: If user is not the uploader.
        """
        attachment = await self.get(attachment_id)
        if not attachment:
            return False

        if attachment.uploaded_by != user_id:
            raise PermissionError("Only the uploader can delete this attachment")

        # Delete GCS files
        try:
            await self.storage.delete(attachment.gcs_path)
        except Exception:
            logger.warning(f"Failed to delete GCS file: {attachment.gcs_path}")
        if attachment.extracted_text_gcs_path:
            try:
                await self.storage.delete(attachment.extracted_text_gcs_path)
            except Exception:
                logger.warning(
                    f"Failed to delete GCS text file: {attachment.extracted_text_gcs_path}"
                )

        # Delete Firestore record
        self.firestore.client.collection(ATTACHMENTS_COLLECTION).document(attachment_id).delete()

        logger.info(f"Deleted attachment {attachment_id}: {attachment.filename}")
        return True

    async def delete_by_session(self, session_id: str) -> int:
        """Delete all attachments belonging to a session.

        Args:
            session_id: Session ID whose attachments should be deleted.

        Returns:
            Number of attachments deleted.
        """
        query = self.firestore.client.collection(ATTACHMENTS_COLLECTION).where(
            "session_id", "==", session_id
        )
        deleted_count = 0
        for doc in query.stream():
            attachment = Attachment.from_firestore(doc.id, doc.to_dict())
            try:
                await self.storage.delete(attachment.gcs_path)
            except Exception:
                logger.warning(f"Failed to delete GCS file: {attachment.gcs_path}")
            if attachment.extracted_text_gcs_path:
                try:
                    await self.storage.delete(attachment.extracted_text_gcs_path)
                except Exception:
                    logger.warning(
                        f"Failed to delete GCS text: {attachment.extracted_text_gcs_path}"
                    )
            doc.reference.delete()
            deleted_count += 1

        if deleted_count:
            logger.info(f"Deleted {deleted_count} attachments for expired session {session_id}")
        return deleted_count
